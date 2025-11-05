# -*- coding: utf-8 -*-
"""
导出渲染器 - 专业试卷导出功能
"""

import os
import uuid
import shutil
import base64
import requests
from datetime import datetime
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import re
from config import LATEX_COMPILE_CONFIG, LATEX_TEMPLATE_PATH, LATEX_CLASS_PATH, LATEX_OUTPUT_DIR

TYPE_SEQUENCE = ["选择题", "填空题", "解答题"]
DEFAULT_QUESTION_TYPE = "解答题"
TYPE_SETTINGS = {
    "选择题": {
        "display": "选择题",
        "score": 5,
        "summary": "{display}，每题{score}分，共{total}分。每个题目的多个选项中只有一个是正确的。"
    },
    "填空题": {
        "display": "填空题",
        "score": 5,
        "summary": "{display}，每题{score}分，共{total}分。请将答案填写在指定位置，不要写出推导过程。"
    },
    "解答题": {
        "display": "解答题",
        "score": 10,
        "summary": "{display}，每题{score}分，共{total}分。请写出完整的解题思路和必要的理由。"
    }
}
try:
    from latex2mathml.converter import convert as latex_to_mathml
    LATEX2MATHML_AVAILABLE = True
except ImportError:
    LATEX2MATHML_AVAILABLE = False

class ExportRenderer:
    """导出渲染器类"""
    
    def __init__(self, upload_folder: str = "uploads"):
        """初始化导出渲染器。

        参数:
            upload_folder: 上传文件夹路径。

        返回:
            None。
        """
        self.upload_folder = upload_folder
    
    def render_latex(self, questions: List[Dict], mode: str, title: str, return_metadata: bool = False):
        """生成 LaTeX 格式的试卷内容。

        参数:
            questions: 题目列表。
            mode: 导出模式（`questions` 或 `with-answers`）。
            title: 试卷标题。
            return_metadata: 是否返回额外的文件元数据。

        返回:
            当 `return_metadata` 为 False 时返回 LaTeX 字符串；否则返回 `(latex_content, metadata)` 元组，其中 metadata 包含输出目录及依赖文件信息。
        """
        if not os.path.exists(LATEX_TEMPLATE_PATH):
            raise FileNotFoundError(f"未找到LaTeX模板文件: {LATEX_TEMPLATE_PATH}")

        if not os.path.exists(LATEX_CLASS_PATH):
            raise FileNotFoundError(f"未找到exam-zh.cls文件: {LATEX_CLASS_PATH}")

        os.makedirs(LATEX_OUTPUT_DIR, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir_name = f"paper_{timestamp}_{uuid.uuid4().hex[:6]}"
        output_dir = os.path.join(LATEX_OUTPUT_DIR, output_dir_name)
        os.makedirs(output_dir, exist_ok=True)

        class_target = os.path.join(output_dir, os.path.basename(LATEX_CLASS_PATH))
        shutil.copy2(LATEX_CLASS_PATH, class_target)

        with open(LATEX_TEMPLATE_PATH, 'r', encoding='utf-8') as f:
            template = f.read()

        copied_images: Dict[str, str] = {}
        grouped_questions = self._group_questions_by_type(questions)
        question_sections: List[str] = []
        subject = "数学试卷"
        current_index = 1

        for question_type, type_questions in grouped_questions:
            if not type_questions:
                continue

            section_title = self._format_section_title(question_type, len(type_questions))
            blocks: List[str] = []

            for question in type_questions:
                block = self._build_question_block(
                    question=question,
                    index=current_index,
                    mode=mode,
                    output_dir=output_dir,
                    copied_images=copied_images
                )
                if block:
                    blocks.append(block)
                    current_index += 1

            if blocks:
                section_content = "\n\n".join(blocks)
                question_sections.append(f"\\section{{{section_title}}}\n\n{section_content}")

        if not question_sections:
            question_sections.append("\\section{题目}\\textit{（暂无题目）}")

        question_sections_str = "\n\n".join(question_sections)

        latex_content = (template
                         .replace('{{TITLE}}', title or '数学试卷')
                         .replace('{{SUBJECT}}', subject)
                         .replace('{{QUESTION_SECTIONS}}', question_sections_str))

        safe_title = re.sub(r'[^0-9A-Za-z\u4e00-\u9fa5_-]+', '_', title or 'exam')
        output_tex_path = os.path.join(output_dir, f"{safe_title}.tex")
        with open(output_tex_path, 'w', encoding='utf-8') as f:
            f.write(latex_content)

        if return_metadata:
            # 构建图片文件映射（文件名 -> 完整路径）
            image_files = {}
            for orig_path, filename in copied_images.items():
                image_full_path = os.path.join(output_dir, filename)
                if os.path.exists(image_full_path):
                    image_files[filename] = image_full_path
            
            metadata = {
                'output_dir': output_dir,
                'cls_file': class_target,
                'tex_file': output_tex_path,
                'image_files': image_files
            }
            return latex_content, metadata
        
        return latex_content
    
    def render_docx(self, questions: List[Dict], mode: str, title: str) -> str:
        """生成 Word 格式的试卷文件。

        参数:
            questions: 题目列表。
            mode: 导出模式（`questions` 或 `with-answers`）。
            title: 试卷标题。

        返回:
            生成的 DOCX 文件路径字符串。
        """
        doc = Document()

        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.9)
            section.right_margin = Inches(0.9)

        title_text = title or '数学试卷'
        title_para = doc.add_heading(title_text, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        current_date = datetime.now().strftime("%Y年%m月%d日")
        date_para = doc.add_paragraph(f"日期：{current_date}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        info_run = info_para.add_run("姓名：__________    学号：__________    班级：__________")
        info_run.font.size = Pt(11)
        info_para.paragraph_format.space_after = Pt(12)

        notice_heading = doc.add_paragraph("注意事项")
        notice_heading.style = 'Heading 2'
        notice_items = [
            "答卷前，考生务必将自己的姓名、考生号、考场号和座位号填写答题卡上，用2B铅笔将试卷类型（B）填涂在答题卡相应位置上，将条形码横贴在答题卡右上角“条形码粘贴处”。",
            "作答选择题时，选出每小题答案后，用2B铅笔在答题卡上对应题目选项的答案信息点涂黑；如需改动，用橡皮擦干净后，再选涂其他答案。答案不能答在试卷上。写在试卷、草稿纸和答题卡上的非答题区域均无效。",
            "非选择题必须用黑色字迹的钢笔或签字笔作答，答案必须写在答题卡各题目指定区域内相应位置上；如需改动，先划掉原来的答案，然后再写上新答案；不准使用铅笔和涂改液。不按以上要求作答无效。",
            "考生必须保持答题卡的整洁。考试结束后，将试卷和答题卡一并交回。"
        ]
        for item in notice_items:
            notice_para = doc.add_paragraph(item, style='List Number')
            notice_para.paragraph_format.space_after = Pt(4)

        doc.add_paragraph()

        grouped_questions = self._group_questions_by_type(questions)
        current_index = 1

        for question_type, type_questions in grouped_questions:
            if not type_questions:
                continue

            summary_text = self._format_section_title(question_type, len(type_questions))
            section_heading = doc.add_heading(summary_text, level=1)
            section_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            section_heading.paragraph_format.space_before = Pt(12)
            section_heading.paragraph_format.space_after = Pt(6)

            for question in type_questions:
                question_heading = doc.add_paragraph(f"第 {current_index} 题")
                question_heading.style = 'Heading 3'
                question_heading.paragraph_format.space_after = Pt(2)

                latex_content = question.get('latex_content', '') or ''
                if not latex_content.strip():
                    latex_content = '（本题内容暂缺）'

                resolved_images = []
                for img_path in question.get('image') or []:
                    resolved = self._resolve_docx_image_path(img_path)
                    if resolved:
                        resolved_images.append(resolved)

                if resolved_images:
                    table = doc.add_table(rows=1, cols=2)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.autofit = False
                    table.columns[0].width = Inches(5.0)
                    table.columns[1].width = Inches(2.5)
                    self._remove_table_borders(table)

                    left_cell = table.cell(0, 0)
                    right_cell = table.cell(0, 1)
                    left_cell.text = ''
                    right_cell.text = ''
                    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

                    # 使用新的方法处理包含公式的内容
                    self._append_latex_content(left_cell, latex_content)
                    self._add_images_to_cell(right_cell, resolved_images)
                else:
                    # 使用新的方法处理包含公式的内容
                    self._append_latex_content(doc, latex_content)

                reference_answer = question.get('reference_answer', '') or ''
                if mode == 'with-answers' and reference_answer.strip():
                    answer_heading = doc.add_paragraph("参考解答")
                    if answer_heading.runs:
                        answer_heading.runs[0].bold = True
                    answer_heading.paragraph_format.space_before = Pt(6)
                    answer_heading.paragraph_format.space_after = Pt(3)

                    answer_table = doc.add_table(rows=1, cols=1)
                    answer_table.style = 'Table Grid'
                    answer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    answer_cell = answer_table.cell(0, 0)
                    answer_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                    answer_cell.text = ''

                    # 使用新的方法处理包含公式的答案内容
                    if not reference_answer.strip():
                        reference_answer = '（本题内容暂缺）'
                    self._append_latex_content(answer_cell, reference_answer)

                doc.add_paragraph()
                current_index += 1

        filename = f'paper_{uuid.uuid4().hex[:8]}.docx'
        file_path = os.path.join(self.upload_folder, filename)
        doc.save(file_path)

        return file_path

    def render_pdf(self, questions: List[Dict], mode: str, title: str) -> str:
        """生成 PDF 格式试卷（通过 LaTeX 编译服务）。

        参数:
            questions: 题目列表。
            mode: 导出模式（`questions` 或 `with-answers`）。
            title: 试卷标题。

        返回:
            生成的 PDF 文件路径字符串。
        """
        # 首先生成LaTeX内容和元数据
        latex_content, metadata = self.render_latex(questions, mode, title, return_metadata=True)
        
        # 调用LaTeX编译API
        api_url = LATEX_COMPILE_CONFIG["api_url"]
        compile_recipe = LATEX_COMPILE_CONFIG.get("compile_recipe")
        
        # 读取图像文件内容并转换为base64编码
        image_data = {}
        for filename, image_path in metadata['image_files'].items():
            if os.path.exists(image_path):
                with open(image_path, 'rb') as f:
                    image_bytes = f.read()
                    image_base64 = base64.b64encode(image_bytes).decode('utf-8')
                    image_data[filename] = image_base64
        
        # 构建依赖文件信息（只传递图像数据，不传递路径）
        dependencies = {
            'image_files': image_data  # 文件名 -> base64编码的数据
        }
        
        payload = {
            "latex_content": latex_content,
            "dependencies": dependencies
        }
        if compile_recipe:
            payload["compile_recipe"] = compile_recipe
        
        response = requests.post(api_url, json=payload, timeout=120)
        
        if response.status_code == 200:
            result = response.json()
            if result.get("success") and result.get("pdf_base64"):
                # 解码PDF并保存
                pdf_filename = f'paper_{uuid.uuid4().hex[:8]}.pdf'
                pdf_path = os.path.join(self.upload_folder, pdf_filename)
                
                pdf_data = base64.b64decode(result["pdf_base64"])
                with open(pdf_path, 'wb') as f:
                    f.write(pdf_data)
                
                return pdf_path
            else:
                raise Exception(result.get("error", "Unknown error"))
        else:
            raise Exception(f"API request failed: {response.status_code}")

    def _build_question_block(self, question: Dict, index: int, mode: str, output_dir: str,
                              copied_images: Dict[str, str]) -> str:
        """构建单道题目的 LaTeX 片段。

        参数:
            question: 题目信息字典。
            index: 题目序号。
            mode: 导出模式。
            output_dir: 输出目录路径。
            copied_images: 已复制图片的映射表。

        返回:
            对应题目的 LaTeX 文本。
        """
        latex_body = question.get('latex_content', '') or ''
        latex_body = self._convert_enumerate_to_choices(latex_body)

        question_body = latex_body.strip()

        images = question.get('image') or []
        image_filenames = []
        for img_path in images:
            filename = self._copy_image_to_output(img_path, output_dir, copied_images, index)
            if filename:
                image_filenames.append(filename)

        if image_filenames:
            image_block = "\\par\n".join(
                f"\\includegraphics[width=0.32\\textwidth]{{{name}}}" for name in image_filenames
            )
            question_text = question_body or '\\textit{（本题内容暂缺）}'
            latex_body = (
                "\\textfigure[text-width=\\columnwidth,fig-pos=bottom-flushright]{%\n"
                f"{question_text}\n"
                "}{%\n"
                f"{image_block}\n"
                "}"
            )

        if not latex_body.strip():
            latex_body = question_body or '\\textit{（本题内容暂缺）}'


        parts = [
            f"% Question {index}",
            "\\begin{question}",
            latex_body.strip(),
            "\\end{question}"
        ]

        reference_answer = question.get('reference_answer', '') or ''
        if mode == 'with-answers' and reference_answer.strip():
            parts.extend([
                "\\begin{tcolorbox}[colback=white,colframe=black,boxrule=0.8pt]",
                "\\begin{solution}",
                reference_answer.strip(),
                "\\end{solution}",
                "\\end{tcolorbox}"
            ])

        return "\n".join(part for part in parts if part)

    def _normalize_question_type(self, question: Dict) -> str:
        """将题目类型标准化为受支持的类别。

        参数:
            question: 题目信息字典。

        返回:
            标准化后的题目类型字符串。
        """
        q_type = (question.get('question_type') or '').strip()
        if q_type not in TYPE_SETTINGS:
            q_type = DEFAULT_QUESTION_TYPE
        return q_type

    def _group_questions_by_type(self, questions: List[Dict]) -> List[tuple]:
        """按照配置顺序对题目按类型分组。

        参数:
            questions: 题目列表。

        返回:
            包含 (题型, 题目列表) 的元组列表。
        """
        grouped = {question_type: [] for question_type in TYPE_SEQUENCE}
        for question in questions:
            normalized = self._normalize_question_type(question)
            grouped.setdefault(normalized, [])
            grouped[normalized].append(question)
        return [(question_type, grouped.get(question_type, [])) for question_type in TYPE_SEQUENCE]

    def _format_section_title(self, question_type: str, count: int) -> str:
        """根据题型生成分组标题。

        参数:
            question_type: 题型名称。
            count: 该题型题目数量。

        返回:
            格式化后的分组标题字符串。
        """
        settings = TYPE_SETTINGS.get(question_type, TYPE_SETTINGS[DEFAULT_QUESTION_TYPE])
        score = settings.get('score', 0)
        total = score * count if score else count
        return settings['summary'].format(
            display=settings.get('display', question_type),
            score=score,
            total=total,
            count=count
        )

    def _convert_enumerate_to_choices(self, content: str) -> str:
        """将 enumerate 环境转换为 choices 环境。

        参数:
            content: 原始 LaTeX 文本。

        返回:
            转换后的文本。
        """
        if not content:
            return ''
        processed = re.sub(r'\\begin\{enumerate\}(\[[^\]]*\])?', r'\\begin{choices}\1', content)
        processed = re.sub(r'\\end\{enumerate\}', r'\\end{choices}', processed)
        processed = re.sub(r'\\begin\{Enumerate\}(\[[^\]]*\])?', r'\\begin{choices}\1', processed)
        processed = re.sub(r'\\end\{Enumerate\}', r'\\end{choices}', processed)
        return processed

    def _copy_image_to_output(self, image_path: str, output_dir: str,
                                 copied_images: Dict[str, str], index: int) -> str:
        """复制图片到输出目录并返回新文件名。

        参数:
            image_path: 原始图片路径。
            output_dir: 导出目录。
            copied_images: 已复制图片映射。
            index: 题目序号。

        返回:
            新图片文件名，若复制失败返回空字符串。
        """
        if not image_path:
            return ''

        if image_path in copied_images:
            return copied_images[image_path]

        clean_path = image_path.lstrip('/')
        if clean_path.startswith('uploads/'):
            clean_path = clean_path[len('uploads/'):]

        absolute_path = os.path.join(self.upload_folder, clean_path)
        if not os.path.exists(absolute_path):
            print(f"图片文件不存在: {absolute_path}")
            return ''

        base_name = os.path.basename(clean_path)
        name, ext = os.path.splitext(base_name)
        dest_filename = f"q{index}_{name}{ext}"
        dest_path = os.path.join(output_dir, dest_filename)
        counter = 1
        while os.path.exists(dest_path):
            dest_filename = f"q{index}_{name}_{counter}{ext}"
            dest_path = os.path.join(output_dir, dest_filename)
            counter += 1

        shutil.copy2(absolute_path, dest_path)
        copied_name = dest_filename.replace('\\', '/')
        copied_images[image_path] = copied_name
        return copied_name
    
    def _clean_latex_content(self, content: str) -> str:
        """
        清理LaTeX内容，确保格式正确
        
        参数:
            content: 原始 LaTeX 内容。
            
        返回:
            清理后的 LaTeX 内容字符串。
        """
        if not content:
            return ""
        
        # 移除多余的空白行
        content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)
        
        # 确保数学公式正确
        content = re.sub(r'\$([^$]+)\$', r'$\1$', content)
        
        # 处理enumerate环境
        content = re.sub(r'\\begin\{enumerate\}', r'\\begin{enumerate}', content)
        content = re.sub(r'\\end\{enumerate\}', r'\\end{enumerate}', content)
        
        # 处理itemize环境
        content = re.sub(r'\\begin\{itemize\}', r'\\begin{itemize}', content)
        content = re.sub(r'\\end\{itemize\}', r'\\end{itemize}', content)
        
        return content.strip()
    
    def _parse_latex_content(self, latex_content: str) -> List[Dict]:
        """
        解析LaTeX内容，将其分割为文本片段和公式片段
        
        参数:
            latex_content: LaTeX 内容字符串。
            
        返回:
            包含 {'type': 'text'/'formula'/'display_formula', 'content': ...} 的列表。
        """
        if not latex_content:
            return []
        
        parts = []
        last_pos = 0
        i = 0
        content_len = len(latex_content)
        
        while i < content_len:
            # 先尝试匹配块级公式 \[...\]
            if latex_content[i:].startswith('\\['):
                # 找到对应的结束标记
                end_pos = latex_content.find('\\]', i + 2)
                if end_pos != -1:
                    # 添加前面的文本
                    before = latex_content[last_pos:i].strip()
                    if before:
                        parts.append({'type': 'text', 'content': before})
                    
                    # 添加公式
                    formula = latex_content[i + 2:end_pos].strip()
                    parts.append({'type': 'display_formula', 'content': formula})
                    
                    last_pos = end_pos + 2
                    i = last_pos
                    continue
            
            # 尝试匹配块级公式 $$...$$
            if latex_content[i:].startswith('$$'):
                # 找到对应的结束标记
                end_pos = latex_content.find('$$', i + 2)
                if end_pos != -1:
                    # 添加前面的文本
                    before = latex_content[last_pos:i].strip()
                    if before:
                        parts.append({'type': 'text', 'content': before})
                    
                    # 添加公式
                    formula = latex_content[i + 2:end_pos].strip()
                    parts.append({'type': 'display_formula', 'content': formula})
                    
                    last_pos = end_pos + 2
                    i = last_pos
                    continue
            
            # 尝试匹配内联公式 $...$（但不匹配 $$）
            if latex_content[i] == '$' and (i + 1 >= content_len or latex_content[i + 1] != '$'):
                # 找到对应的结束 $
                end_pos = latex_content.find('$', i + 1)
                if end_pos != -1:
                    # 添加前面的文本
                    before = latex_content[last_pos:i].strip()
                    if before:
                        parts.append({'type': 'text', 'content': before})
                    
                    # 添加公式
                    formula = latex_content[i + 1:end_pos].strip()
                    parts.append({'type': 'formula', 'content': formula})
                    
                    last_pos = end_pos + 1
                    i = last_pos
                    continue
            
            i += 1
        
        # 添加剩余的文本
        if last_pos < content_len:
            remaining = latex_content[last_pos:].strip()
            if remaining:
                parts.append({'type': 'text', 'content': remaining})
        
        return parts if parts else [{'type': 'text', 'content': latex_content}]
    
    def _add_latex_formula_to_run(self, run, latex_formula: str, is_display: bool = False):
        """
        将LaTeX公式添加到Word运行对象中（使用Office Math格式）
        
        参数:
            run: Word运行对象。
            latex_formula: LaTeX公式字符串。
            is_display: 是否为块级公式。
        """
        if not LATEX2MATHML_AVAILABLE:
            # 如果latex2mathml不可用，添加原始文本
            run.add_text(f" [{latex_formula}] ")
            return
        
        try:
            # 转换LaTeX为MathML
            mathml_str = latex_to_mathml(latex_formula)
            
            # 将MathML转换为Office Math XML
            omath = self._mathml_to_omath_element(mathml_str)
            
            if omath is not None:
                # 将oMath添加到运行元素
                run._element.append(omath)
            else:
                # 如果转换失败，使用简单文本
                run.add_text(f" [{latex_formula}] ")
                
        except Exception as e:
            print(f"LaTeX公式转换失败 ({latex_formula[:50] if len(latex_formula) > 50 else latex_formula}...): {e}")
            # 如果转换失败，添加原始文本
            run.add_text(f" [{latex_formula}] ")
    
    def _mathml_to_omath_element(self, mathml_str: str):
        """
        将MathML字符串转换为Office Math XML元素
        
        参数:
            mathml_str: MathML字符串。
            
        返回:
            Office Math XML元素，如果转换失败返回None。
        """
        try:
            import xml.etree.ElementTree as ET
            import traceback
            
            # 清理MathML字符串
            mathml_clean = mathml_str.strip()
            
            # 移除XML声明和命名空间声明（如果存在）
            if mathml_clean.startswith('<?xml'):
                mathml_clean = mathml_clean.split('>', 1)[1] if '>' in mathml_clean else mathml_clean
            
            # 尝试解析MathML
            mathml_root = None
            parse_error = None
            try:
                mathml_root = ET.fromstring(mathml_clean)
            except ET.ParseError as e:
                parse_error = e
                # 如果解析失败，尝试添加math标签包装
                try:
                    mathml_wrapped = f'<math xmlns="http://www.w3.org/1998/Math/MathML">{mathml_clean}</math>'
                    mathml_root = ET.fromstring(mathml_wrapped)
                except Exception as e2:
                    print(f"MathML解析失败 - 原始MathML: {mathml_clean[:200]}...")
                    print(f"第一次解析错误: {parse_error}")
                    print(f"第二次解析错误: {e2}")
                    print(f"错误类型: {type(e2).__name__}")
                    import traceback
                    traceback.print_exc()
                    return None
            
            # 创建Office Math元素
            # 使用parse_xml创建包含命名空间的元素
            omath_xml = '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"></m:oMath>'
            omath = parse_xml(omath_xml)
            
            # 递归转换MathML元素到Office Math
            def convert_mathml_to_omath(mathml_elem, omath_parent):
                """递归转换MathML元素"""
                try:
                    # 安全地获取标签名
                    tag_name = mathml_elem.tag
                    if '}' in tag_name:
                        tag_name = tag_name.split('}')[-1]
                    
                    if tag_name == 'math' or tag_name == 'mrow':
                        # 处理子元素
                        for child in mathml_elem:
                            convert_mathml_to_omath(child, omath_parent)
                        # 处理文本内容
                        if mathml_elem.text and mathml_elem.text.strip():
                            r = OxmlElement('m:r')
                            t = OxmlElement('m:t')
                            t.text = mathml_elem.text.strip()
                            r.append(t)
                            omath_parent.append(r)
                    
                    elif tag_name == 'mi' or tag_name == 'mn' or tag_name == 'mo':
                        # 数学标识符、数字、运算符
                        r = OxmlElement('m:r')
                        t = OxmlElement('m:t')
                        text = mathml_elem.text or ''
                        if tag_name == 'mi':
                            # 标识符使用斜体
                            rpr = OxmlElement('m:rPr')
                            i = OxmlElement('m:i')
                            i.set(qn('m:val'), '1')
                            rpr.append(i)
                            r.append(rpr)
                        if text:
                            t.text = text.strip()
                            r.append(t)
                            omath_parent.append(r)
                        
                        # 处理子元素
                        for child in mathml_elem:
                            convert_mathml_to_omath(child, omath_parent)
                    
                    elif tag_name == 'mfrac':
                        # 分数
                        f = OxmlElement('m:f')
                        num = OxmlElement('m:num')
                        den = OxmlElement('m:den')
                        child_count = 0
                        for child in mathml_elem:
                            child_tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                            if child_tag == 'mrow':
                                for subchild in child:
                                    convert_mathml_to_omath(subchild, num if child_count == 0 else den)
                            else:
                                convert_mathml_to_omath(child, num if child_count == 0 else den)
                            child_count += 1
                        f.append(num)
                        f.append(den)
                        omath_parent.append(f)
                    
                    elif tag_name == 'msup' or tag_name == 'msub' or tag_name == 'msubsup':
                        # 上标、下标、上下标
                        if tag_name == 'msup':
                            func = OxmlElement('m:sSup')
                            e = OxmlElement('m:e')
                            sup = OxmlElement('m:sup')
                        elif tag_name == 'msub':
                            func = OxmlElement('m:sSub')
                            e = OxmlElement('m:e')
                            sub = OxmlElement('m:sub')
                        else:  # msubsup
                            func = OxmlElement('m:sSubSup')
                            e = OxmlElement('m:e')
                            sub = OxmlElement('m:sub')
                            sup = OxmlElement('m:sup')
                        
                        children = list(mathml_elem)
                        if children:
                            convert_mathml_to_omath(children[0], e)
                            func.append(e)
                            if len(children) > 1:
                                convert_mathml_to_omath(children[1], sub if tag_name == 'msub' else sup)
                                func.append(sub if tag_name == 'msub' else sup)
                            if len(children) > 2:  # msubsup
                                convert_mathml_to_omath(children[2], sup)
                                func.append(sup)
                        omath_parent.append(func)
                    
                    elif tag_name == 'msqrt' or tag_name == 'mroot':
                        # 根号
                        if tag_name == 'msqrt':
                            rad = OxmlElement('m:rad')
                            deg = None
                        else:
                            rad = OxmlElement('m:rad')
                            deg = OxmlElement('m:deg')
                            children = list(mathml_elem)
                            if len(children) > 1:
                                convert_mathml_to_omath(children[1], deg)
                        
                        e = OxmlElement('m:e')
                        children = list(mathml_elem)
                        if children:
                            convert_mathml_to_omath(children[0], e)
                        rad.append(e)
                        if deg:
                            rad.append(deg)
                        omath_parent.append(rad)
                    
                    else:
                        # 其他元素，递归处理
                        for child in mathml_elem:
                            convert_mathml_to_omath(child, omath_parent)
                        if mathml_elem.text and mathml_elem.text.strip():
                            r = OxmlElement('m:r')
                            t = OxmlElement('m:t')
                            t.text = mathml_elem.text.strip()
                            r.append(t)
                            omath_parent.append(r)
                            
                except Exception as inner_e:
                    print(f"转换MathML元素时出错 - 标签: {getattr(mathml_elem, 'tag', 'unknown')}")
                    print(f"错误类型: {type(inner_e).__name__}")
                    print(f"错误信息: {inner_e}")
                    import traceback
                    traceback.print_exc()
                    raise
            
            # 开始转换
            convert_mathml_to_omath(mathml_root, omath)
            
            # 如果转换后没有内容，返回None
            if len(omath) == 0:
                return None
            
            return omath
            
        except Exception as e:
            print(f"MathML转换失败")
            print(f"错误类型: {type(e).__name__}")
            print(f"错误信息: {e}")
            print(f"原始MathML (前500字符): {mathml_str[:500] if mathml_str else 'None'}")
            import traceback
            traceback.print_exc()
            return None
    
    def _create_simple_omath_from_latex(self, latex_or_mathml: str):
        """
        从LaTeX或MathML创建简单的Office Math元素（后备方案）
        
        参数:
            latex_or_mathml: LaTeX公式或MathML字符串。
            
        返回:
            Office Math XML元素。
        """
        # 使用parse_xml创建包含命名空间的元素
        omath_xml = '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"></m:oMath>'
        omath = parse_xml(omath_xml)
        
        # 创建一个简单的文本运行
        r = OxmlElement('m:r')
        t = OxmlElement('m:t')
        # 只取前100个字符，避免过长
        display_text = latex_or_mathml[:100] if len(latex_or_mathml) > 100 else latex_or_mathml
        t.text = display_text
        r.append(t)
        omath.append(r)
        
        return omath
    
    def _latex_to_readable(self, latex_content: str) -> str:
        """
        将LaTeX内容转换为可读文本（用于Word文档，保留公式标记以便后续处理）
        
        参数:
            latex_content: LaTeX 内容字符串。
            
        返回:
            转换后的可读文本（公式已标记）。
        """
        if not latex_content:
            return ""
        
        # 处理列表环境，转换为文本格式
        content = latex_content
        
        # 处理enumerate环境
        content = re.sub(r'\\begin\{enumerate\}', '', content)
        content = re.sub(r'\\end\{enumerate\}', '', content)
        content = re.sub(r'\\item\s*', '• ', content)
        
        # 处理itemize环境
        content = re.sub(r'\\begin\{itemize\}', '', content)
        content = re.sub(r'\\end\{itemize\}', '', content)
        
        # 处理其他常见的LaTeX命令（但不处理数学公式）
        # 移除文本格式命令，保留公式
        content = re.sub(r'\\(?:textbf|textit|text|emph)\{([^}]*)\}', r'\1', content)
        content = re.sub(r'\\(?:textbf|textit|text|emph)\s+([^\s]+)', r'\1', content)
        
        # 清理多余空白
        content = re.sub(r'[ \t]+', ' ', content)
        content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)
        
        return content.strip()

    def _resolve_docx_image_path(self, image_path: str) -> Optional[str]:
        """将图片路径解析为可用于 DOCX 的绝对路径。

        参数:
            image_path: 图片路径字符串。

        返回:
            可用的绝对路径；若不可用则返回 None。
        """
        if not image_path:
            return None

        clean_path = image_path.strip()
        if not clean_path:
            return None

        if clean_path.startswith('http://') or clean_path.startswith('https://'):
            return None  # 当前不支持远程图片

        if clean_path.startswith('/'):
            clean_path = clean_path.lstrip('/')

        if clean_path.startswith('uploads/'):
            candidate = os.path.join(self.upload_folder, clean_path[len('uploads/'):])
        else:
            candidate = clean_path if os.path.isabs(clean_path) else os.path.join(self.upload_folder, clean_path)

        if os.path.exists(candidate):
            return candidate
        return None

    def _append_latex_content(self, container, latex_content: str, bullet_style: str = 'List Bullet') -> None:
        """
        向 docx 容器追加包含LaTeX公式的内容
        
        参数:
            container: docx 段落或单元格对象。
            latex_content: 包含LaTeX公式的内容字符串。
            bullet_style: 项目符号样式名称。
            
        返回:
            None。
        """
        if not latex_content:
            return
        
        # 解析LaTeX内容，分割为文本和公式片段
        parts = self._parse_latex_content(latex_content)
        
        if not parts:
            return
        
        # 按行处理（处理换行）
        lines = []
        current_line_parts = []
        
        for part in parts:
            if part['type'] == 'text':
                # 检查文本中是否包含换行
                text_lines = part['content'].split('\n')
                if len(text_lines) > 1:
                    # 如果有换行，先保存当前行的内容
                    if current_line_parts:
                        lines.append(current_line_parts)
                        current_line_parts = []
                    # 处理多行文本
                    for i, text_line in enumerate(text_lines):
                        if i > 0:
                            # 新行
                            if current_line_parts:
                                lines.append(current_line_parts)
                                current_line_parts = []
                        if text_line.strip():
                            current_line_parts.append({'type': 'text', 'content': text_line})
                else:
                    # 单行文本
                    if text_lines[0].strip():
                        current_line_parts.append(part)
            else:
                # 公式片段
                current_line_parts.append(part)
        
        # 添加最后一行
        if current_line_parts:
            lines.append(current_line_parts)
        
        # 如果没有行，添加一个空行
        if not lines:
            lines = [[]]
        
        # 写入段落
        is_table_cell = hasattr(container, '_tc')
        existing_paragraphs = list(getattr(container, 'paragraphs', [])) if is_table_cell else []
        first_paragraph = existing_paragraphs[0] if existing_paragraphs else None
        
        for idx, line_parts in enumerate(lines):
            if idx == 0 and first_paragraph is not None:
                paragraph = first_paragraph
                paragraph.text = ''
            else:
                paragraph = container.add_paragraph()
            
            # 检查是否是列表项
            is_bullet = False
            if line_parts and line_parts[0]['type'] == 'text':
                text = line_parts[0]['content']
                if text.startswith('•'):
                    is_bullet = True
                    line_parts[0]['content'] = text[1:].strip()
                    if bullet_style:
                        paragraph.style = bullet_style
            
            # 添加内容片段
            for part in line_parts:
                if part['type'] == 'text':
                    text = part['content'].strip()
                    if text:
                        run = paragraph.add_run(text)
                elif part['type'] == 'formula':
                    # 内联公式
                    run = paragraph.add_run()
                    self._add_latex_formula_to_run(run, part['content'], is_display=False)
                elif part['type'] == 'display_formula':
                    # 块级公式（居中显示）
                    run = paragraph.add_run()
                    self._add_latex_formula_to_run(run, part['content'], is_display=True)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            paragraph.paragraph_format.space_after = Pt(6)
    
    def _append_text_lines(self, container, lines: List[str], bullet_style: str = 'List Bullet') -> None:
        """向 docx 容器追加文本行并处理项目符号。

        参数:
            container: docx 段落或单元格对象。
            lines: 待写入的文本行列表。
            bullet_style: 项目符号样式名称。

        返回:
            None。
        """
        is_table_cell = hasattr(container, '_tc')
        existing_paragraphs = list(getattr(container, 'paragraphs', [])) if is_table_cell else []
        first_paragraph = existing_paragraphs[0] if existing_paragraphs else None

        if not lines:
            lines = ['']

        for idx, raw_line in enumerate(lines):
            line = raw_line or ''
            is_bullet = line.startswith('•')
            text = line[1:].strip() if is_bullet else line.strip()

            if idx == 0 and first_paragraph is not None:
                paragraph = first_paragraph
                paragraph.text = ''
            else:
                paragraph = container.add_paragraph()

            if is_bullet and bullet_style:
                paragraph.style = bullet_style
            run = paragraph.add_run(text or ' ')
            paragraph.paragraph_format.space_after = Pt(6)

    def _add_images_to_cell(self, cell, image_paths: List[str]) -> None:
        """按顺序将图片插入到表格单元格内。

        参数:
            cell: docx 单元格对象。
            image_paths: 图片路径列表。

        返回:
            None。
        """
        existing_paragraphs = list(getattr(cell, 'paragraphs', []))
        first_paragraph = existing_paragraphs[0] if existing_paragraphs else None

        for idx, image_path in enumerate(image_paths):
            if not os.path.exists(image_path):
                continue

            if idx == 0 and first_paragraph is not None:
                paragraph = first_paragraph
                paragraph.text = ''
            else:
                paragraph = cell.add_paragraph()

            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run()
            try:
                run.add_picture(image_path, width=Inches(2.2))
            except Exception as exc:
                print(f"无法添加图片 {image_path}: {exc}")
            paragraph.paragraph_format.space_after = Pt(6)

    def _remove_table_borders(self, table) -> None:
        """移除 docx 表格的所有边框。

        参数:
            table: docx 表格对象。

        返回:
            None。
        """
        tbl = table._tbl
        
        # 获取或创建 tblPr 元素
        # 使用 find 方法查找，如果不存在则创建
        tbl_pr = tbl.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            # 将 tblPr 插入到表格元素的开始位置
            tbl.insert(0, tbl_pr)
        
        # 查找并移除现有的 tblBorders
        tbl_borders = tbl_pr.find(qn('w:tblBorders'))
        if tbl_borders is not None:
            tbl_pr.remove(tbl_borders)

        # 创建新的无边框设置
        tbl_borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), 'nil')
            tbl_borders.append(edge_element)
        tbl_pr.append(tbl_borders)
    
    def _add_mathml_to_paragraph(self, paragraph, latex_math: str):
        """将 LaTeX 数学公式转换为 MathML 并插入段落。

        参数:
            paragraph: Word 段落对象。
            latex_math: LaTeX 数学公式字符串。

        返回:
            None。
        """
        if not LATEX2MATHML_AVAILABLE or not latex_math:
            return
        
        try:
            # 转换LaTeX为MathML
            mathml = latex_to_mathml(latex_math)
            
            # 创建MathML元素，使用parse_xml创建包含命名空间的元素
            math_element_xml = '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"></m:oMath>'
            math_element = parse_xml(math_element_xml)
            
            # 解析MathML并添加到元素中
            mathml_xml = parse_xml(f'<math xmlns="http://www.w3.org/1998/Math/MathML">{mathml}</math>')
            math_element.append(mathml_xml)
            
            # 添加到段落
            paragraph._element.append(math_element)
            
        except Exception as e:
            print(f"MathML转换失败: {e}")
            # 如果转换失败，添加原始文本
            paragraph.add_run(latex_math)
