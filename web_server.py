# -*- coding: utf-8 -*-
"""
高考题目录入和自动打标系统Web服务器
"""

from flask import Flask, render_template, request, jsonify, send_from_directory, session, redirect, url_for
from functools import wraps
import json
import os
import uuid
from werkzeug.utils import secure_filename
from question_manager import QuestionManager
from ocr_client import DeepSeekOCRClient
from user_manager import UserManager
from config import WEB_CONFIG, QUESTION_TAGS, OCR_BASE_URL, LLM_CONFIG, SECRET_KEY, USER_DATABASE_PATH

app = Flask(__name__)
app.secret_key = SECRET_KEY

# 配置上传文件夹
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# 确保上传文件夹存在
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# 初始化题目管理器
question_manager = QuestionManager()

# 初始化OCR客户端
ocr_client = DeepSeekOCRClient(OCR_BASE_URL)

# 初始化用户管理器
user_manager = UserManager(USER_DATABASE_PATH)

# 登录验证装饰器
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login_page'))
        return f(*args, **kwargs)
    return decorated_function

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
@login_required
def index():
    """主页"""
    user = user_manager.get_user_by_id(session['user_id'])
    return render_template('index.html', user=user)

@app.route('/login')
def login_page():
    """登录页面"""
    return render_template('login.html')

@app.route('/api/auth/register', methods=['POST'])
def register():
    """用户注册API"""
    try:
        data = request.get_json()
        username = data.get('username')
        password = data.get('password')
        email = data.get('email')
        
        success, message = user_manager.register_user(username, password, email)
        
        return jsonify({
            'success': success,
            'message': message
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/auth/login', methods=['POST'])
def login():
    """用户登录API"""
    try:
        data = request.get_json()
        username = data.get('username')
        password = data.get('password')
        
        user = user_manager.authenticate_user(username, password)
        
        if user:
            session['user_id'] = user['id']
            session['username'] = user['username']
            return jsonify({
                'success': True,
                'message': '登录成功',
                'user': user
            })
        else:
            return jsonify({
                'success': False,
                'message': '用户名或密码错误'
            })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/auth/logout', methods=['POST'])
def logout():
    """用户登出API"""
    session.clear()
    return jsonify({'success': True, 'message': '已登出'})

@app.route('/api/auth/current', methods=['GET'])
@login_required
def get_current_user():
    """获取当前登录用户信息API"""
    user = user_manager.get_user_by_id(session['user_id'])
    if user:
        return jsonify({
            'success': True,
            'user': user
        })
    return jsonify({'success': False, 'message': '未登录'}), 401

@app.route('/static/<path:filename>')
def static_files(filename):
    """静态文件服务"""
    return send_from_directory('static', filename)

@app.route('/api/questions', methods=['POST'])
@login_required
def add_question():
    """添加题目API"""
    try:
        data = request.get_json()
        
        # 验证必需字段
        if not data or 'latex_content' not in data:
            return jsonify({'success': False, 'message': '题目内容不能为空'}), 400
        
        latex_content = data['latex_content']
        tags = data.get('tags', [])
        reference_answer = data.get('reference_answer', '')
        source = data.get('source', '')
        image = data.get('image', [])
        visibility = data.get('visibility', 'public')  # 新增：可见范围，默认所有人
        
        # 添加题目
        question_id = question_manager.add_question(
            latex_content=latex_content,
            tags=tags,
            reference_answer=reference_answer,
            source=source,
            image=image,
            user_id=session['user_id'],
            visibility=visibility
        )
        
        return jsonify({
            'success': True, 
            'message': '题目添加成功',
            'question_id': question_id
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/questions/auto-tag', methods=['POST'])
@login_required
def auto_tag_question():
    """自动打标和生成解答API"""
    try:
        data = request.get_json()
        
        if not data or 'latex_content' not in data:
            return jsonify({'success': False, 'message': '题目内容不能为空'}), 400
        
        latex_content = data['latex_content']
        source = data.get('source', '')
        
        # 自动打标和生成解答
        tags, answer = question_manager.auto_tag_and_answer(latex_content, source)
        
        return jsonify({
            'success': True,
            'tags': tags,
            'answer': answer
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/questions/search', methods=['GET'])
@login_required
def search_questions():
    """搜索题目API"""
    try:
        # 获取查询参数
        tags = request.args.getlist('tags')
        keyword = request.args.get('keyword', '')
        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 20))
        offset = (page - 1) * limit
        
        # 获取当前用户ID
        current_user_id = session['user_id']
        
        questions = []
        
        if tags:
            # 按标签查询
            questions = question_manager.get_questions_by_tags(tags, current_user_id)
        elif keyword:
            # 关键词搜索
            questions = question_manager.search_questions(keyword, current_user_id)
        else:
            # 获取所有题目
            questions = question_manager.get_all_questions(limit, offset, current_user_id)
        
        return jsonify({
            'success': True,
            'questions': questions,
            'total': len(questions)
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/questions/<int:question_id>', methods=['GET'])
@login_required
def get_question(question_id):
    """获取单个题目详情API"""
    try:
        current_user_id = session['user_id']
        question = question_manager.get_question_by_id(question_id, current_user_id)
        
        if question:
            return jsonify({
                'success': True,
                'question': question
            })
        else:
            return jsonify({
                'success': False,
                'message': '题目不存在或无权访问'
            }), 404
            
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/questions/<int:question_id>/answer', methods=['GET'])
@login_required
def get_question_answer(question_id):
    """获取题目参考解答API"""
    try:
        current_user_id = session['user_id']
        answer = question_manager.get_reference_answer(question_id, current_user_id)
        
        if answer:
            return jsonify({
                'success': True,
                'answer': answer
            })
        else:
            return jsonify({
                'success': False,
                'message': '参考解答不存在或无权访问'
            }), 404
            
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/questions/stats', methods=['GET'])
@login_required
def get_question_stats():
    """获取题目统计信息API"""
    try:
        current_user_id = session['user_id']
        stats = question_manager.get_question_stats(current_user_id)
        return jsonify({
            'success': True,
            'stats': stats
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/upload', methods=['POST'])
@login_required
def upload_file():
    """上传图片API"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': '没有选择文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'message': '没有选择文件'}), 400
        
        if file and allowed_file(file.filename):
            # 生成唯一文件名
            filename = secure_filename(file.filename)
            name, ext = os.path.splitext(filename)
            unique_filename = f"{uuid.uuid4()}{ext}"
            
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
            file.save(file_path)
            
            return jsonify({
                'success': True,
                'filename': unique_filename,
                'url': f'/uploads/{unique_filename}'
            })
        else:
            return jsonify({'success': False, 'message': '不支持的文件类型'}), 400
            
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    """提供上传的图片文件"""
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/api/questions/<int:question_id>', methods=['DELETE'])
@login_required
def delete_question(question_id):
    """删除题目API"""
    try:
        current_user_id = session['user_id']
        success = question_manager.delete_question(question_id, current_user_id)
        
        if success:
            return jsonify({
                'success': True,
                'message': '题目删除成功'
            })
        else:
            return jsonify({
                'success': False,
                'message': '题目不存在或无权删除'
            }), 404
            
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/ocr-parse', methods=['POST'])
@login_required
def ocr_parse():
    """OCR解析试卷API"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': '没有选择文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'message': '没有选择文件'}), 400
        
        if file and allowed_file(file.filename):
            # 保存临时文件
            filename = secure_filename(file.filename)
            name, ext = os.path.splitext(filename)
            temp_filename = f"temp_{uuid.uuid4()}{ext}"
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], temp_filename)
            file.save(temp_path)
            
            try:
                # 调用OCR
                ocr_result = ocr_client.ocr_image(temp_path)
                markdown_content = ocr_result.get('markdown', '')
                
                # 调用大模型解析题目
                parsed_questions = question_manager.parse_exam_paper(markdown_content)
                
                return jsonify({
                    'success': True,
                    'questions': parsed_questions
                })
                
            except Exception as e:
                return jsonify({'success': False, 'message': f'OCR解析失败: {str(e)}'}), 500
            finally:
                # 删除临时文件
                if os.path.exists(temp_path):
                    os.remove(temp_path)
        else:
            return jsonify({'success': False, 'message': '不支持的文件类型'}), 400
            
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/tags', methods=['GET'])
def get_tags():
    """获取所有可用标签API"""
    return jsonify({
        'success': True,
        'tags': QUESTION_TAGS
    })

@app.route('/api/export-paper', methods=['POST'])
@login_required
def export_paper():
    """导出试卷API"""
    try:
        data = request.get_json()
        questions = data.get('questions', [])
        mode = data.get('mode', 'questions')  # questions 或 with-answers
        format_type = data.get('format', 'latex')  # latex, docx, 或 pdf
        
        if not questions:
            return jsonify({'success': False, 'message': '没有题目可导出'}), 400
        
        # 生成文件内容
        if format_type == 'latex':
            content = generate_latex_paper(questions, mode)
            mimetype = 'text/plain'
            filename = f'paper_{uuid.uuid4().hex[:8]}.tex'
        elif format_type == 'docx':
            file_path = generate_docx_paper(questions, mode)
            return send_from_directory(os.path.dirname(file_path), os.path.basename(file_path), 
                                     as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        elif format_type == 'pdf':
            file_path = generate_pdf_paper(questions, mode)
            return send_from_directory(os.path.dirname(file_path), os.path.basename(file_path), 
                                     as_attachment=True, mimetype='application/pdf')
        else:
            return jsonify({'success': False, 'message': '不支持的格式'}), 400
        
        # 对于LaTeX，直接返回内容
        from flask import Response
        return Response(
            content,
            mimetype=mimetype,
            headers={"Content-disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

def generate_latex_paper(questions, mode):
    """生成LaTeX格式试卷"""
    content = r"""\documentclass[12pt,a4paper]{article}
\usepackage{ctex}
\usepackage{amsmath}
\usepackage{amssymb}
\usepackage{geometry}
\geometry{left=2.5cm,right=2.5cm,top=2.5cm,bottom=2.5cm}

\title{试卷}
\author{}
\date{\today}

\begin{document}
\maketitle

"""
    
    # 添加题目
    for i, q in enumerate(questions, 1):
        content += f"\\section*{{题目 {i}}}\n\n"
        content += q['latex_content'] + "\n\n"
        
        if mode == 'with-answers' and q.get('reference_answer'):
            content += "\\subsection*{参考解答}\n\n"
            content += q['reference_answer'] + "\n\n"
    
    content += r"\end{document}"
    return content

def generate_docx_paper(questions, mode):
    """生成Word格式试卷"""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # 标题
    title = doc.add_heading('试卷', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加题目
    for i, q in enumerate(questions, 1):
        doc.add_heading(f'题目 {i}', level=1)
        
        # 题目内容
        p = doc.add_paragraph(q['latex_content'])
        
        # 如果包含答案
        if mode == 'with-answers' and q.get('reference_answer'):
            doc.add_heading('参考解答', level=2)
            doc.add_paragraph(q['reference_answer'])
        
        # 添加分隔
        if i < len(questions):
            doc.add_paragraph()
    
    # 保存文件
    filename = f'paper_{uuid.uuid4().hex[:8]}.docx'
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    doc.save(file_path)
    
    return file_path

def generate_pdf_paper(questions, mode):
    """生成PDF格式试卷"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    
    filename = f'paper_{uuid.uuid4().hex[:8]}.pdf'
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    
    doc = SimpleDocTemplate(file_path, pagesize=A4)
    story = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        alignment=1
    )
    
    # 标题
    story.append(Paragraph('试卷', title_style))
    story.append(Spacer(1, 0.5 * inch))
    
    # 添加题目
    for i, q in enumerate(questions, 1):
        story.append(Paragraph(f'<b>题目 {i}</b>', styles['Heading2']))
        story.append(Spacer(1, 0.2 * inch))
        
        # 题目内容（简化处理LaTeX）
        content = q['latex_content'].replace('$', '').replace('\\', '')
        story.append(Paragraph(content, styles['Normal']))
        
        # 如果包含答案
        if mode == 'with-answers' and q.get('reference_answer'):
            story.append(Spacer(1, 0.2 * inch))
            story.append(Paragraph('<b>参考解答</b>', styles['Heading3']))
            answer = q['reference_answer'].replace('$', '').replace('\\', '')
            story.append(Paragraph(answer, styles['Normal']))
        
        story.append(Spacer(1, 0.5 * inch))
    
    doc.build(story)
    return file_path

@app.errorhandler(404)
def not_found(error):
    """404错误处理"""
    return jsonify({'success': False, 'message': '页面不存在'}), 404

@app.errorhandler(500)
def internal_error(error):
    """500错误处理"""
    return jsonify({'success': False, 'message': '服务器内部错误'}), 500

if __name__ == '__main__':
    # 确保templates目录存在
    if not os.path.exists('templates'):
        os.makedirs('templates')
    
    print(f"启动Web服务器: http://{WEB_CONFIG['host']}:{WEB_CONFIG['port']}")
    app.run(
        host=WEB_CONFIG['host'],
        port=WEB_CONFIG['port'],
        debug=WEB_CONFIG['debug']
    )
